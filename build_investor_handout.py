from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

OUT = "output/handout/CEAtlas_Investor_Budget_Handout.docx"

NAVY = "17324D"
BLUE = "2878B5"
PALE_BLUE = "EAF2F8"
LIGHT = "F3F5F7"
MID = "647383"
INK = "18222D"
GREEN = "2E7D5B"
PALE_GREEN = "EAF5EF"
WHITE = "FFFFFF"
GOLD = "A66F00"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=85, start=110, bottom=85, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + side
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=10.5, bold=False, color=INK, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", size=10.5, bold=False, color=INK, after=5, before=0,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_with_next = keep
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(label.upper())
    set_run(r, size=8.5, bold=True, color=label_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    p2.paragraph_format.line_spacing = 1.08
    r2 = p2.add_run(text)
    set_run(r2, size=11, bold=True, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=8.5, color=MID)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.10

for name, size, before, after, color in [
    ("Heading 1", 16, 10, 6, NAVY),
    ("Heading 2", 12.5, 8, 4, BLUE),
    ("Heading 3", 11, 6, 3, NAVY),
]:
    st = styles[name]
    st.font.name = "Arial"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

list_style = styles["List Bullet"]
list_style.font.name = "Arial"
list_style.font.size = Pt(10.5)
list_style.paragraph_format.left_indent = Inches(0.33)
list_style.paragraph_format.first_line_indent = Inches(-0.17)
list_style.paragraph_format.space_after = Pt(3)

section.different_first_page_header_footer = False

def populate_header_footer(header, footer):
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("CEATLAS  |  CONFIDENTIAL DISCUSSION DRAFT")
    set_run(hr, size=8.2, bold=True, color=MID)
    fp = footer.paragraphs[0]
    fp.text = ""
    fr = fp.add_run("CEAtlas - $500,000 Pre-Seed Use of Funds")
    set_run(fr, size=8.5, color=MID)
    if len(footer.paragraphs) > 1:
        pp = footer.paragraphs[1]
        pp.text = ""
    else:
        pp = footer.add_paragraph()
    add_page_number(pp)

for hdr, ftr in [
    (section.header, section.footer),
    (section.first_page_header, section.first_page_footer),
    (section.even_page_header, section.even_page_footer),
]:
    populate_header_footer(hdr, ftr)

# PAGE 1
add_paragraph(doc, "PRE-SEED INVESTMENT BRIEF", size=9, bold=True, color=BLUE, after=4)
add_paragraph(doc, "CEAtlas", size=28, bold=True, color=NAVY, after=1)
add_paragraph(doc, "Turning a functioning CE discovery beta into a reliable transaction-enabled marketplace", size=13.5, color=MID, after=12)

add_callout(doc, "The ask", "$500,000 for approximately 18 months of runway, targeting roughly 8-9% ownership on final agreed terms.", fill=PALE_GREEN, label_color=GREEN)

add_heading(doc, "What already exists", 1)
add_paragraph(doc, "CEAtlas is a live beta for dental continuing-education discovery and planning. It is not a concept-only pitch: the product already contains a structured catalog, provider data, search and comparison workflows, accounts, maps, state-requirement guidance, provider connections, and travel-related functionality.", after=6)

metric = doc.add_table(rows=1, cols=3)
metric.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(metric, [3120, 3120, 3120], indent=120)
for idx, (value, label) in enumerate([
    ("8,677", "courses and events displayed"),
    ("167", "providers and hosts displayed"),
    ("1", "working marketplace foundation"),
]):
    cell = metric.cell(0, idx)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(value)
    set_run(r, size=19, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(1)
    r2 = p2.add_run(label)
    set_run(r2, size=8.5, color=MID)

add_heading(doc, "What the financing changes", 1)
for text, lead in [
    ("Founder full-time. The founder leaves outside employment and focuses on product, providers, distribution, and commercialization.", "Founder full-time."),
    ("Production-grade engineering. A staged senior engineer audits, stabilizes, tests, and extends the existing platform before the company commits to a larger team.", "Production-grade engineering."),
    ("Provider-side commerce. CEAtlas builds toward provider claiming, course management, analytics, direct registration, and transaction revenue.", "Provider-side commerce."),
    ("Evidence for the next decision. The company uses the runway to measure acquisition, provider activation, conversion, retention, and marketplace revenue.", "Evidence for the next decision."),
]:
    add_bullet(doc, text, lead)

add_callout(doc, "Investment rationale", "The capital is not being raised to start building CEAtlas. It is being raised to convert an unusually capital-efficient product head start into a reliable, defensible, revenue-producing company.")

doc.add_page_break()

# PAGE 2
add_paragraph(doc, "USE OF FUNDS", size=9, bold=True, color=BLUE, after=4)
add_paragraph(doc, "$500,000 - 18-Month Operating Plan", size=22, bold=True, color=NAVY, after=2)
add_paragraph(doc, "$421,200 of planned operating investment plus a $78,800 controlled reserve.", size=11.5, color=MID, after=10)

rows = [
    ("Founder compensation", "$135,000", "27.0%", "$90,000 annual gross salary for 18 months"),
    ("Staged engineering", "$145,000", "29.0%", "Audit, stabilization, launch reliability, provider tools and transactions"),
    ("Founder payroll burden", "$16,200", "3.2%", "Employer payroll taxes and benefit allowance; engineering assumes contractor staging"),
    ("Security / penetration testing", "$15,000", "3.0%", "Independent web/API review, written findings and remediation retest"),
    ("Cloud / data / software", "$20,000", "4.0%", "Hosting, database, scraping compute, monitoring, analytics and operating software"),
    ("Product / UI design", "$12,000", "2.4%", "UX review, launch polish, accessibility and provider/checkout workflows"),
    ("Marketing / user acquisition", "$30,000", "6.0%", "SEO/content and measured channel experiments after analytics is reliable"),
    ("Provider acquisition / sales", "$12,000", "2.4%", "Founder-led outreach, CRM, onboarding assets and integration support"),
    ("Legal / corporate / IP", "$18,000", "3.6%", "Formation, financing, founder stock, IP assignment, contracts and policies"),
    ("Accounting / payroll / tax", "$8,000", "1.6%", "Bookkeeping, payroll, tax filings and investor-ready reporting"),
    ("Insurance / administration", "$5,000", "1.0%", "Cyber/E&O/D&O/general coverage, registered agent and state administration"),
    ("Travel / industry research", "$5,000", "1.0%", "Selective provider meetings, conferences and travel-market validation"),
    ("Unallocated runway reserve", "$78,800", "15.8%", "Overruns, an unexpectedly productive channel, added engineering or runway extension"),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Table Grid"
headers = ["Category", "Amount", "%", "Purpose"]
for i, h in enumerate(headers):
    c = table.rows[0].cells[i]
    set_cell_shading(c, NAVY)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = c.paragraphs[0]
    r = p.add_run(h)
    set_run(r, size=8.5, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for ridx, row in enumerate(rows):
    cells = table.add_row().cells
    if ridx % 2 == 1:
        for c in cells:
            set_cell_shading(c, LIGHT)
    for idx, value in enumerate(row):
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(value)
        set_run(r, size=7.9 if idx == 3 else 8.3, bold=(idx in (0, 1)), color=INK)
set_table_geometry(table, [2200, 1250, 700, 5210], indent=120)

total = doc.add_table(rows=1, cols=3)
total.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(total, [5200, 2080, 2080], indent=120)
for c in total.rows[0].cells:
    set_cell_shading(c, PALE_GREEN)
vals = [("TOTAL", True), ("$500,000", True), ("100.0%", True)]
for c, (txt, bold) in zip(total.rows[0].cells, vals):
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c is not total.rows[0].cells[0] else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(txt)
    set_run(r, size=10.5, bold=bold, color=GREEN)

add_paragraph(doc, "Budget discipline: engineering is staged, security follows initial remediation, marketing is released through measured experiments, and the reserve is not treated as automatic spending.", size=9.2, italic=True, color=MID, before=6, after=0)

doc.add_page_break()

# PAGE 3
add_paragraph(doc, "MILESTONES AND CAPITAL JUSTIFICATION", size=9, bold=True, color=BLUE, after=4)
add_paragraph(doc, "What $500,000 Is Expected to Produce", size=22, bold=True, color=NAVY, after=8)

milestones = [
    ("Months 0-3", "Foundation", "Form company and banking; founder full-time; engineering audit; restore catalog pipeline; fix major defects; install analytics; complete first security remediation."),
    ("Months 3-6", "Prove demand", "Relaunch a polished core experience; improve catalog quality; begin provider claiming; conduct founder-led provider interviews; run controlled acquisition tests."),
    ("Months 6-12", "Build the marketplace", "Provider dashboard and course-management workflows; transaction MVP; user accounts and CE planning; initial monetization; measure conversion and retention."),
    ("Months 12-18", "Prove the company", "Scale the channels that work; expand provider acquisition and professions; validate travel monetization; grow transaction activity; prepare the next financing or sustainability plan."),
]
mt = doc.add_table(rows=1, cols=3)
mt.style = "Table Grid"
mt.alignment = WD_TABLE_ALIGNMENT.LEFT
for idx, text in enumerate(["Timing", "Objective", "Expected progress"]):
    c = mt.rows[0].cells[idx]
    set_cell_shading(c, NAVY)
    r = c.paragraphs[0].add_run(text)
    set_run(r, size=9, bold=True, color=WHITE)
set_repeat_table_header(mt.rows[0])
for ridx, row in enumerate(milestones):
    cells = mt.add_row().cells
    if ridx % 2 == 1:
        for c in cells:
            set_cell_shading(c, LIGHT)
    for idx, text in enumerate(row):
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(text)
        set_run(r, size=9, bold=(idx < 2), color=INK)
set_table_geometry(mt, [1450, 1800, 6110], indent=120)

add_heading(doc, "Why the amount is justified", 1)
for text, lead in [
    ("Runway protects execution. An 18-month plan reduces the risk that the company must fundraise again before engineering and commercial evidence materially improve.", "Runway protects execution."),
    ("The product has a real head start. Capital is applied to reliability, transactions and distribution rather than beginning from a blank page.", "The product has a real head start."),
    ("Engineering risk is controlled. The budget supports a senior audit and staged engagement before any permanent technical-team decision.", "Engineering risk is controlled."),
    ("Commercial learning stays founder-led. Early provider conversations remain with the founder until a repeatable value proposition and sales process emerge.", "Commercial learning stays founder-led."),
    ("The reserve is strategic, not decorative. It allows the company to respond to the actual bottleneck while maintaining a cash floor.", "The reserve is strategic, not decorative."),
]:
    add_bullet(doc, text, lead)

add_heading(doc, "Capital controls", 1)
controls = doc.add_table(rows=2, cols=3)
controls.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(controls, [3120, 3120, 3120], indent=120)
control_items = [
    ("MONTHLY", "Reforecast burn, runway and reserve"),
    ("STAGE GATES", "Approve engineering, security and growth in phases"),
    ("MONTH 12", "Begin next-financing or sustainability decision"),
]
for idx, (label, text) in enumerate(control_items):
    top = controls.rows[0].cells[idx]
    bot = controls.rows[1].cells[idx]
    set_cell_shading(top, PALE_BLUE)
    set_cell_shading(bot, LIGHT)
    p = top.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    set_run(r, size=8.3, bold=True, color=BLUE)
    p2 = bot.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(text)
    set_run(r2, size=9, bold=True, color=NAVY)

add_callout(doc, "The outcome", "By month 18, CEAtlas should be able to demonstrate whether it can become a scalable marketplace: reliable supply, activated providers, measurable user demand, transaction behavior, and a credible path to revenue growth.", fill=PALE_GREEN, label_color=GREEN)

add_paragraph(doc, "Discussion draft - figures are planning estimates and remain subject to vendor quotes, legal structure, worker classification, benefits and final financing terms.", size=8.2, italic=True, color=MID, before=4, after=0)

doc.core_properties.title = "CEAtlas Investor Budget Handout"
doc.core_properties.subject = "$500,000 pre-seed use of funds and 18-month milestone plan"
doc.core_properties.author = "CEAtlas"
doc.save(OUT)
print(OUT)
